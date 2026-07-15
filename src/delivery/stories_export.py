"""Delivery-backlog exports: a styled Word document and an RFC 4180 CSV.

Both exports consume the validated story package produced by
``src.delivery.models`` so every row and heading is already traceable to
requirements and transcript utterances.
"""
from __future__ import annotations

import csv
import io
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.delivery.docx_export import (
    MUTED,
    _add_label_value_table,
    _add_table,
    _body,
    _clean,
    _configure_document,
    _evidence,
    _heading,
    _page_field,
    _run,
)
from src.intelligence.state import Utterance

CSV_HEADER = ["Issue Type", "Key", "Summary", "Description",
              "Acceptance Criteria", "Epic", "Requirements", "Evidence"]
# ID | Given | When | Then; widths must sum to the shared 9360 DXA content width.
AC_TABLE_WIDTHS = [1100, 2760, 2740, 2760]


def build_stories_docx(package: dict[str, Any], meta: dict[str, Any]) -> DocumentObject:
    """Build a delivery-backlog Word document in the BRD's visual language.

    ``meta`` is the stored session dict (``SessionStore.load_session``): its
    utterances supply the mm:ss evidence timestamps and its meta block supplies
    the session identity shown in the document header.
    """
    utterances: list[Utterance] = meta.get("utterances", [])
    times = {u.id: u.t0 for u in utterances}
    epics = [item for item in package.get("epics", []) if isinstance(item, dict)]
    stories = [item for item in package.get("stories", []) if isinstance(item, dict)]

    doc = Document()
    _configure_document(doc)
    _add_header_footer(doc, str(meta.get("id", "")))

    title = _clean(package.get("title")) or "Delivery Backlog"
    paragraph = doc.add_paragraph(style="Title")
    paragraph.add_run(title)
    paragraph = doc.add_paragraph(style="Subtitle")
    paragraph.add_run("Epics and user stories with acceptance criteria and transcript traceability")

    _add_label_value_table(doc, [
        ("Session", str(meta.get("id", ""))),
        ("Started", str(meta.get("meta", {}).get("started", "Not recorded"))),
        ("Backlog revision", str(package.get("revision", 0))),
        ("Scope", f"{len(epics)} epic{'s' if len(epics) != 1 else ''} · "
                  f"{len(stories)} user stor{'ies' if len(stories) != 1 else 'y'}"),
    ])

    if not epics:
        _body(doc, "No epics or stories have been generated for this session.", italic=True)
    for index, epic in enumerate(epics, 1):
        _heading(doc, f"{index}. {epic.get('id', '')} — {_clean(epic.get('title')) or 'Untitled epic'}", 1)
        if _clean(epic.get("description")):
            _body(doc, _clean(epic.get("description")))
        _traceability(doc, epic, times)
        for story in stories:
            if story.get("epic_id") != epic.get("id"):
                continue
            _heading(doc, f"{story.get('id', '')} — {_clean(story.get('title')) or 'Untitled story'}", 2)
            _body(doc, _narrative(story))
            criteria = [item for item in story.get("acceptance_criteria", []) if isinstance(item, dict)]
            if criteria:
                rows = [[str(criterion.get("id", "")), _clean(criterion.get("given")),
                         _clean(criterion.get("when")), _clean(criterion.get("then"))]
                        for criterion in criteria]
                _add_table(doc, ["ID", "Given", "When", "Then"], rows, AC_TABLE_WIDTHS)
            else:
                _body(doc, "No acceptance criteria were captured.", italic=True)
            _traceability(doc, story, times)

    doc.core_properties.title = title
    doc.core_properties.subject = "ReqPilot delivery backlog"
    doc.core_properties.author = "ReqPilot"
    return doc


def stories_csv(package: dict[str, Any]) -> str:
    """Flatten the package into an import-friendly CSV (RFC 4180, CRLF)."""
    epics = [item for item in package.get("epics", []) if isinstance(item, dict)]
    stories = [item for item in package.get("stories", []) if isinstance(item, dict)]
    epic_titles = {str(epic.get("id", "")): _clean(epic.get("title")) for epic in epics}

    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\r\n")
    writer.writerow(CSV_HEADER)
    for epic in epics:
        writer.writerow([
            "Epic", str(epic.get("id", "")), _clean(epic.get("title")),
            _clean(epic.get("description")), "", "",
            _joined(epic.get("requirement_ids")), _evidence_ids(epic),
        ])
    for story in stories:
        epic_id = str(story.get("epic_id", ""))
        epic_label = f"{epic_id}: {epic_titles[epic_id]}" if epic_titles.get(epic_id) else epic_id
        writer.writerow([
            "Story", str(story.get("id", "")), _clean(story.get("title")),
            _narrative(story), _criteria_sentences(story), epic_label,
            _joined(story.get("requirement_ids")), _evidence_ids(story),
        ])
    return buffer.getvalue()


def _narrative(story: dict[str, Any]) -> str:
    return (f"As a {_clean(story.get('as_a'))}, I want {_clean(story.get('i_want'))}, "
            f"so that {_clean(story.get('so_that'))}.")


def _criteria_sentences(story: dict[str, Any]) -> str:
    sentences = [
        f"Given {_clean(item.get('given'))} When {_clean(item.get('when'))} Then {_clean(item.get('then'))}"
        for item in story.get("acceptance_criteria", []) if isinstance(item, dict)
    ]
    return " | ".join(sentences)


def _joined(values: Any) -> str:
    return ", ".join(str(value) for value in values) if isinstance(values, list) else ""


def _evidence_ids(item: dict[str, Any]) -> str:
    refs = item.get("evidence_utterances")
    return ", ".join(f"U{uid}" for uid in refs) if isinstance(refs, list) else ""


def _traceability(doc: DocumentObject, item: dict[str, Any], times: dict[int, float]) -> None:
    requirements = _joined(item.get("requirement_ids")) or "Not linked"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    _run(paragraph.add_run(
        f"Traceability: requirements {requirements} · evidence {_evidence(item, times)}"
    ), 9, MUTED)


def _add_header_footer(doc: DocumentObject, session_id: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    _run(header.add_run("REQPILOT  |  DELIVERY BACKLOG"), 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    _run(footer.add_run(f"Session {session_id}   |   Page "), 8.5, MUTED)
    _page_field(footer)
