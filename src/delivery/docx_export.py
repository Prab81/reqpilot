"""Polished, deterministic BRD DOCX export using standard_business_brief tokens."""
from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.intelligence import prompts
from src.intelligence.state import Utterance, mmss
from src.sessions.store import SessionStore

if TYPE_CHECKING:
    from src.intelligence.providers import LlmProvider


CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5B6573"
LIGHT_FILL = "F2F4F7"
BORDER = "C9D1DA"


def export_brd_docx(session_id: str, store: SessionStore, provider: "LlmProvider",
                    output_path: Path | str) -> Path:
    """Export the stored session as a styled Word BRD.

    Narrative prose is LLM-assisted; requirements, decisions, gaps, and all
    transcript evidence are assembled from stored data so traceability cannot
    be lost when the provider fails or omits content.
    """
    session = store.load_session(session_id)
    system, user = prompts.build_brd_prompt(session["state"], session["utterances"])
    from src.intelligence.providers import ProviderError

    try:
        narrative = provider.complete_json(system, user, prompts.BRD_SCHEMA_HINT, max_tokens=8192)
    except ProviderError:
        narrative = {}
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    build_brd_docx(session, path, narrative)
    return path


def build_brd_docx(session: dict[str, Any], output_path: Path | str,
                   narrative: dict[str, Any] | None = None) -> Path:
    narrative = narrative if isinstance(narrative, dict) else {}
    state = session.get("state", {})
    utterances: list[Utterance] = session.get("utterances", [])
    times = {u.id: u.t0 for u in utterances}
    doc = Document()
    _configure_document(doc)
    _add_header_footer(doc, str(session.get("id", "")))

    title = _clean(narrative.get("title")) or "Business Requirements Document"
    subtitle = _clean(state.get("title")) or _clean(session.get("meta", {}).get("title"))
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    if subtitle and subtitle.casefold() not in title.casefold():
        p = doc.add_paragraph(style="Subtitle")
        p.add_run(subtitle)

    duration = max((u.t1 for u in utterances), default=0.0)
    meta = [
        ("Session", str(session.get("id", ""))),
        ("Started", str(session.get("meta", {}).get("started", "Not recorded"))),
        ("Duration", mmss(duration)),
        ("Evidence", f"{len(utterances)} timestamped transcript segments"),
    ]
    _add_label_value_table(doc, meta)

    _heading(doc, "1. Executive Summary", 1)
    summary = state.get("summary", [])
    if isinstance(summary, list) and summary:
        for item in summary:
            if _clean(item):
                _list_item(doc, _clean(item))
    else:
        _body(doc, "No executive summary was captured during the session.", italic=True)

    _heading(doc, "2. Context and Background", 1)
    _body(doc, _clean(narrative.get("context")) or _summary_fallback(summary))

    _heading(doc, "3. Stakeholders", 1)
    stakeholders = narrative.get("stakeholders", [])
    rows: list[list[str]] = []
    if isinstance(stakeholders, list):
        for stakeholder in stakeholders:
            if isinstance(stakeholder, dict) and _clean(stakeholder.get("name")):
                rows.append([_clean(stakeholder.get("name")), _clean(stakeholder.get("interest"))])
            elif _clean(stakeholder):
                rows.append([_clean(stakeholder), ""])
    if rows:
        _add_table(doc, ["Stakeholder", "Interest / responsibility"], rows, [2200, 7160])
    else:
        _body(doc, "Stakeholders were not explicitly identified.", italic=True)

    _heading(doc, "4. Process Overview", 1)
    _heading(doc, "Current process", 2)
    _body(doc, _clean(narrative.get("current_process")) or "The current process was not captured in narrative form.")
    _heading(doc, "Target process", 2)
    _body(doc, _clean(narrative.get("future_process")) or "The target process was not captured in narrative form.")
    diagrams = state.get("diagrams", [])
    if isinstance(diagrams, list) and diagrams:
        _heading(doc, "Captured process models", 2)
        for diagram in diagrams:
            if not isinstance(diagram, dict):
                continue
            refs = _evidence(diagram, times)
            _list_item(doc, f"{diagram.get('title', diagram.get('id', 'Process model'))} - evidence {refs}")

    _heading(doc, "5. Functional Requirements", 1)
    requirement_rows = []
    for requirement in state.get("requirements", []):
        requirement_rows.append([
            str(requirement.get("id", "")),
            _clean(requirement.get("text")),
            _clean(requirement.get("status")),
            _evidence(requirement, times),
        ])
    if requirement_rows:
        _add_table(doc, ["ID", "Requirement", "Status", "Transcript evidence"],
                   requirement_rows, [700, 5740, 1120, 1800])
    else:
        _body(doc, "No functional requirements were captured.", italic=True)

    _heading(doc, "6. Non-Functional Requirements", 1)
    nfrs = [item for item in narrative.get("non_functional", []) if _clean(item)] \
        if isinstance(narrative.get("non_functional"), list) else []
    nfr_gaps = [g for g in state.get("gaps", []) if g.get("category") == "nfr"]
    if nfrs or nfr_gaps:
        for item in nfrs:
            _list_item(doc, _clean(item))
        for gap in nfr_gaps:
            _list_item(doc, f"Open NFR gap {gap.get('id', '')}: {_clean(gap.get('text'))} "
                            f"(evidence {_evidence(gap, times)})")
    else:
        _body(doc, "No non-functional requirements were explicitly captured.", italic=True)

    _heading(doc, "7. Decisions and Assumptions", 1)
    _heading(doc, "Decisions", 2)
    decisions = state.get("decisions", [])
    if decisions:
        for decision in decisions:
            _list_item(doc, f"{decision.get('id', '')}: {_clean(decision.get('text'))} "
                            f"(evidence {_evidence(decision, times)})", numbered=True)
    else:
        _body(doc, "No decisions were captured.", italic=True)
    _heading(doc, "Assumptions", 2)
    assumptions = narrative.get("assumptions", [])
    if isinstance(assumptions, list) and assumptions:
        for assumption in assumptions:
            if _clean(assumption):
                _list_item(doc, _clean(assumption))
    else:
        _body(doc, "No assumptions were recorded.", italic=True)

    _heading(doc, "8. Open Items", 1)
    open_rows: list[list[str]] = []
    for question in state.get("open_questions", []):
        if question.get("status") != "answered":
            open_rows.append([
                str(question.get("id", "")),
                str(question.get("status", "suggested")),
                _clean(question.get("category")),
                _clean(question.get("text")),
            ])
    for gap in state.get("gaps", []):
        if gap.get("category") != "nfr":
            open_rows.append([
                str(gap.get("id", "")), "gap", _clean(gap.get("category")),
                f"{_clean(gap.get('text'))} Evidence: {_evidence(gap, times)}",
            ])
    if open_rows:
        _add_table(doc, ["ID", "Status", "Category", "Open question or gap"],
                   open_rows, [650, 1050, 1250, 6410])
    else:
        _body(doc, "No open questions or gaps remain.")

    _heading(doc, "9. Transcript Evidence Register", 1)
    _body(doc, "Session-relative timestamps below are the source of truth for requirement traceability.")
    evidence_rows = [
        [f"U{u.id}", f"{mmss(u.t0)}-{mmss(u.t1)}", u.text]
        for u in utterances
    ]
    if evidence_rows:
        _add_table(doc, ["Segment", "Timestamp", "Transcript"], evidence_rows,
                   [800, 1450, 7110])
    else:
        _body(doc, "No transcript segments were stored.", italic=True)

    doc.core_properties.title = title
    doc.core_properties.subject = subtitle or "ReqPilot business requirements"
    doc.core_properties.author = "ReqPilot"
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def _configure_document(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    _style_font(styles["Normal"], "Calibri", 11, "000000")
    normal = styles["Normal"].paragraph_format
    normal.space_before = Pt(0)
    normal.space_after = Pt(6)
    normal.line_spacing = 1.10

    _style_font(styles["Title"], "Calibri", 24, NAVY, bold=True)
    styles["Title"].paragraph_format.space_before = Pt(8)
    styles["Title"].paragraph_format.space_after = Pt(4)
    styles["Title"].paragraph_format.keep_with_next = True
    _style_font(styles["Subtitle"], "Calibri", 13, MUTED)
    styles["Subtitle"].paragraph_format.space_before = Pt(0)
    styles["Subtitle"].paragraph_format.space_after = Pt(14)
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        _style_font(styles[name], "Calibri", size, color, bold=True)
        fmt = styles[name].paragraph_format
        fmt.space_before, fmt.space_after = Pt(before), Pt(after)
        fmt.keep_with_next = True
        fmt.keep_together = True

    for style_name in ("ReqPilot Bullet", "ReqPilot Number"):
        if style_name not in styles:
            style = styles.add_style(style_name, 1)
        else:
            style = styles[style_name]
        style.base_style = styles["Normal"]
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    bullet_num, number_num = _create_numbering(doc)
    styles["ReqPilot Bullet"]._element.set(qn("w:numId"), str(bullet_num))
    styles["ReqPilot Number"]._element.set(qn("w:numId"), str(number_num))
    doc._reqpilot_num_ids = (bullet_num, number_num)  # type: ignore[attr-defined]


def _style_font(style: Any, name: str, size: float, color: str,
                bold: bool | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def _create_numbering(doc: DocumentObject) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add(abstract_id: int, num_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
        num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), fmt); level.append(num_fmt)
        level_text = OxmlElement("w:lvlText"); level_text.set(qn("w:val"), text); level.append(level_text)
        suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); level.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720")
        tabs.append(tab); ppr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360")
        ppr.append(ind); level.append(ppr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref); numbering.append(num)

    add(next_abstract, next_num, "bullet", "•")
    add(next_abstract + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def _add_header_footer(doc: DocumentObject, session_id: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("REQPILOT  |  BUSINESS REQUIREMENTS DOCUMENT")
    _run(run, 8.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run(f"Session {session_id}   |   Page ")
    _run(run, 8.5, MUTED)
    _page_field(footer)


def _page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t"); value.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    _run(run, 8.5, MUTED)


def _heading(doc: DocumentObject, text: str, level: int) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def _body(doc: DocumentObject, text: str, italic: bool = False) -> None:
    paragraphs = [part.strip() for part in text.replace("\r", "").split("\n\n") if part.strip()]
    for part in paragraphs or [""]:
        p = doc.add_paragraph()
        run = p.add_run(part)
        run.italic = italic


def _list_item(doc: DocumentObject, text: str, numbered: bool = False) -> None:
    style = "ReqPilot Number" if numbered else "ReqPilot Bullet"
    paragraph = doc.add_paragraph(style=style)
    num_id = doc._reqpilot_num_ids[1 if numbered else 0]  # type: ignore[attr-defined]
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    paragraph.add_run(text)


def _add_label_value_table(doc: DocumentObject, values: list[tuple[str, str]]) -> None:
    # Metadata is a definition list rather than comparable row/column data.
    # Paragraphs are more accessible than a headerless layout table and remain
    # easy to scan in Word, LibreOffice, and Google Docs imports.
    for label, value in values:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        _run(label_run, 10, NAVY, bold=True)
        value_run = paragraph.add_run(value)
        _run(value_run, 10, "000000")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def _add_table(doc: DocumentObject, headers: list[str], rows: list[list[str]],
               widths: list[int], *, header: bool = True, quiet: bool = False) -> Any:
    if sum(widths) != CONTENT_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table = doc.add_table(rows=0, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if header and headers:
        row = table.add_row()
        for index, text in enumerate(headers):
            _set_cell_text(row.cells[index], text, bold=True, color=NAVY, align=_column_alignment(index, widths))
            _shade(row.cells[index], LIGHT_FILL)
        _repeat_header(row)
    for row_values in rows:
        row = table.add_row()
        for index, cell in enumerate(row.cells):
            value = str(row_values[index]) if index < len(row_values) else ""
            _set_cell_text(cell, value, bold=quiet and index == 0,
                           color=NAVY if quiet and index == 0 else "000000",
                           align=_column_alignment(index, widths))
    _set_table_geometry(table, widths, quiet=quiet)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def _column_alignment(index: int, widths: list[int]) -> Any:
    return WD_ALIGN_PARAGRAPH.CENTER if widths[index] <= 1800 else WD_ALIGN_PARAGRAPH.LEFT


def _set_cell_text(cell: Any, text: str, *, bold: bool, color: str, align: Any) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    _run(run, 9.3, color, bold=bold)


def _run(run: Any, size: float, color: str, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _set_table_geometry(table: Any, widths: list[int], *, quiet: bool) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblBorders", "w:tblCellMar"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    tbl_w = OxmlElement("w:tblW"); tbl_w.set(qn("w:w"), str(CONTENT_DXA)); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd"); tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA)); tbl_ind.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed")
    tbl_pr.extend([tbl_w, tbl_ind, layout])
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil" if quiet else "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:color"), BORDER)
        borders.append(elem)
    tbl_pr.append(borders)
    margins = OxmlElement("w:tblCellMar")
    for name, value in CELL_MARGINS.items():
        elem = OxmlElement(f"w:{name}"); elem.set(qn("w:w"), str(value)); elem.set(qn("w:type"), "dxa")
        margins.append(elem)
    tbl_pr.append(margins)
    grid = tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader"); marker.set(qn("w:val"), "true"); tr_pr.append(marker)


def _evidence(item: dict[str, Any], times: dict[int, float]) -> str:
    refs = [f"{mmss(times[uid])} (U{uid})" for uid in item.get("evidence_utterances", []) if uid in times]
    return ", ".join(refs) if refs else "Not linked"


def _summary_fallback(summary: Any) -> str:
    if isinstance(summary, list):
        values = [_clean(item) for item in summary if _clean(item)]
        if values:
            return " ".join(values)
    return "Context and background were not captured in narrative form."


def _clean(value: Any) -> str:
    return value.strip() if isinstance(value, str) else ""
