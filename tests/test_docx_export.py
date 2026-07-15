from __future__ import annotations

import base64
import io
import struct
import zlib
from zipfile import ZipFile

from docx import Document
from fastapi.testclient import TestClient
from lxml import etree

from src.delivery import export_brd_docx
from src.intelligence.providers import MockProvider
from src.server import create_app
from src.sessions.store import SessionStore
from tests.fixtures.meeting_transcript import brd_narrative, pass2_state, utterances


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _png_1x1() -> bytes:
    """A minimal, structurally valid 1x1 white RGB PNG."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (struct.pack(">I", len(data)) + tag + data
                + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF))
    header = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    pixels = zlib.compress(b"\x00\xff\xff\xff")
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", header)
            + chunk(b"IDAT", pixels) + chunk(b"IEND", b""))


def _session(tmp_path) -> tuple[SessionStore, str]:
    store = SessionStore(tmp_path / "data")
    sid = store.create_session()
    for utterance in utterances():
        store.append_utterance(sid, utterance)
    store.snapshot_state(sid, pass2_state(), 2, analyzed_through=30)
    return store, sid


def test_docx_export_contains_all_requirements_and_timestamp_evidence(tmp_path) -> None:
    store, sid = _session(tmp_path)
    path = export_brd_docx(sid, store, MockProvider([brd_narrative()]), tmp_path / "BRD.docx")

    document = Document(path)
    text = "\n".join(p.text for p in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "Business Requirements Document" in text or "BRD" in text
    assert "Transcript Evidence Register" in text
    for requirement in pass2_state()["requirements"]:
        assert requirement["id"] in table_text
        assert requirement["text"] in table_text
    assert "00:08 (U2)" in table_text
    assert "U30" in table_text and "04:01-04:08" in table_text


def test_docx_uses_real_styles_numbering_and_fixed_dxa_geometry(tmp_path) -> None:
    store, sid = _session(tmp_path)
    path = export_brd_docx(sid, store, MockProvider([brd_narrative()]), tmp_path / "BRD.docx")

    with ZipFile(path) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        styles_xml = etree.fromstring(archive.read("word/styles.xml"))
        numbering_xml = etree.fromstring(archive.read("word/numbering.xml"))
    assert document_xml.xpath("count(.//w:pStyle[@w:val='Heading1'])", namespaces=NS) >= 1
    assert document_xml.xpath("count(.//w:numPr/w:numId)", namespaces=NS) >= 1
    assert numbering_xml.xpath("count(.//w:abstractNum/w:lvl/w:numFmt[@w:val='bullet'])", namespaces=NS) >= 1
    assert styles_xml.xpath("string(.//w:style[@w:styleId='Normal']/w:pPr/w:spacing/@w:after)", namespaces=NS) == "120"
    for table in document_xml.xpath(".//w:tbl", namespaces=NS):
        assert table.xpath("string(w:tblPr/w:tblW/@w:type)", namespaces=NS) == "dxa"
        assert table.xpath("string(w:tblPr/w:tblW/@w:w)", namespaces=NS) == "9360"
        assert table.xpath("string(w:tblPr/w:tblInd/@w:w)", namespaces=NS) == "120"
        assert table.xpath(
            "count(w:tr[1]/w:trPr/w:tblHeader)", namespaces=NS
        ) == 1
        grid = [int(value) for value in table.xpath("w:tblGrid/w:gridCol/@w:w", namespaces=NS)]
        assert sum(grid) == 9360
        for row in table.xpath("w:tr", namespaces=NS):
            widths = [int(value) for value in row.xpath("w:tc/w:tcPr/w:tcW/@w:w", namespaces=NS)]
            assert widths == grid


def test_docx_degrades_to_deterministic_content_when_provider_fails(tmp_path) -> None:
    store, sid = _session(tmp_path)
    path = export_brd_docx(sid, store, MockProvider([]), tmp_path / "fallback.docx")

    document = Document(path)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert path.stat().st_size > 10_000
    assert "R14" in table_text
    assert "03:53 (U29)" in table_text


def test_docx_embeds_valid_diagram_png_with_caption_and_evidence(tmp_path) -> None:
    store, sid = _session(tmp_path)
    png_base64 = base64.b64encode(_png_1x1()).decode("ascii")

    path = export_brd_docx(sid, store, MockProvider([brd_narrative()]),
                           tmp_path / "with-image.docx", diagram_images={"G1": png_base64})

    document = Document(path)
    assert len(document.inline_shapes) == 1
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Loan application processing flow" in text
    assert "Evidence 00:08 (U2)" in text  # caption evidence line under the image


def test_docx_falls_back_to_text_line_for_invalid_diagram_payloads(tmp_path) -> None:
    store, sid = _session(tmp_path)
    not_png = base64.b64encode(b"plain text, not an image").decode("ascii")

    for name, images in (
        ("bad-base64.docx", {"G1": "not base64 at all!!!"}),
        ("not-png.docx", {"G1": not_png}),
        ("unknown-id.docx", {"G404": base64.b64encode(_png_1x1()).decode("ascii")}),
    ):
        path = export_brd_docx(sid, store, MockProvider([brd_narrative()]),
                               tmp_path / name, diagram_images=images)
        document = Document(path)
        assert len(document.inline_shapes) == 0
        text = "\n".join(p.text for p in document.paragraphs)
        assert "Loan application processing flow - evidence" in text


def test_brd_docx_post_route_embeds_images_and_survives_garbage(tmp_path) -> None:
    store, sid = _session(tmp_path)
    client = TestClient(create_app(store=store, provider=MockProvider([])))
    png_base64 = base64.b64encode(_png_1x1()).decode("ascii")

    with_image = client.post(f"/api/session/{sid}/brd.docx",
                             json={"diagrams": [{"id": "G1", "png_base64": png_base64}]})
    assert with_image.status_code == 200
    assert with_image.headers["content-type"].startswith("application/vnd.openxmlformats")
    document = Document(io.BytesIO(with_image.content))
    assert len(document.inline_shapes) == 1

    garbage = client.post(f"/api/session/{sid}/brd.docx",
                          json={"diagrams": [{"id": "G1", "png_base64": "not base64 at all!!!"}]})
    assert garbage.status_code == 200
    document = Document(io.BytesIO(garbage.content))
    assert len(document.inline_shapes) == 0
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Loan application processing flow - evidence" in text

    # The bare GET contract is unchanged: no body, text-only diagram line.
    plain = client.get(f"/api/session/{sid}/brd.docx")
    assert plain.status_code == 200
    assert len(Document(io.BytesIO(plain.content)).inline_shapes) == 0


def test_brd_docx_post_route_rejects_oversized_image_payloads(tmp_path) -> None:
    store, sid = _session(tmp_path)
    client = TestClient(create_app(store=store, provider=MockProvider([])))

    oversized = "A" * (28 * 1024 * 1024)  # ~21 MB decoded, over the 20 MB cap
    response = client.post(f"/api/session/{sid}/brd.docx",
                           json={"diagrams": [{"id": "G1", "png_base64": oversized}]})
    assert response.status_code == 413
    assert "20 MB" in response.json()["detail"]
