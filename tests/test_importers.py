from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document

from src.importers import MAX_IMPORT_BYTES, TranscriptImportError, import_transcript


FIXTURES = Path(__file__).parent / "fixtures"


def test_teams_vtt_preserves_speakers_and_exact_cue_times() -> None:
    result = import_transcript(
        (FIXTURES / "teams_transcript.vtt").read_bytes(), "teams_transcript.vtt"
    )

    assert result.format == "vtt"
    assert [(u.t0, u.t1) for u in result.utterances] == [
        (2.25, 6.75), (7.1, 11.4), (12.0, 14.25),
    ]
    assert [u.text for u in result.utterances] == [
        "Priya Shah: Applications must be reviewed within twenty-four hours.",
        "Daniel Wu: If a document is missing, notify the applicant.",
        "Priya Shah: The notification should be email and SMS.",
    ]


def test_copilot_pasted_text_recognizes_common_turn_header_shapes() -> None:
    source = (FIXTURES / "copilot_recap.txt").read_text(encoding="utf-8")
    result = import_transcript(source, "copilot-recap.txt")

    assert [u.t0 for u in result.utterances] == [5.0, 12.0, 19.0]
    assert [u.text for u in result.utterances] == [
        "Priya Shah: We need one queue for all applications.",
        "Daniel Wu: The queue must be sorted by application age.",
        "Priya Shah: Only senior underwriters may reassign an application.",
    ]
    assert [u.id for u in result.utterances] == [1, 2, 3]


def test_untimed_txt_gets_deterministic_non_overlapping_synthetic_times() -> None:
    result = import_transcript(
        "Analyst: Welcome to discovery.\n\nStakeholder: We need an audit trail.",
        "notes.TXT",
    )

    first, second = result.utterances
    assert first.t0 == 0.0
    assert first.t1 >= 1.0
    assert second.t0 > first.t1
    assert second.text == "Stakeholder: We need an audit trail."


def test_utf16_txt_is_decoded() -> None:
    result = import_transcript("Alice: Hello\n\nBob: Hi".encode("utf-16"), "meeting.txt")
    assert [u.text for u in result.utterances] == ["Alice: Hello", "Bob: Hi"]


def test_docx_reads_paragraphs_and_tables_in_document_order() -> None:
    document = Document()
    document.add_paragraph("Facilitator 00:01")
    document.add_paragraph("Describe the approval process.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Stakeholder"
    table.cell(0, 1).text = "Managers approve requests above $5,000."
    payload = io.BytesIO()
    document.save(payload)

    result = import_transcript(payload.getvalue(), "Workshop.DOCX")

    assert result.format == "docx"
    assert result.utterances[0].t0 == 1.0
    assert result.utterances[0].text == "Facilitator: Describe the approval process."
    assert result.utterances[1].text == "Stakeholder: Managers approve requests above $5,000."


def test_docx_three_column_transcript_table_preserves_timestamp_and_speaker() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "00:02:15"
    table.cell(0, 1).text = "Avery Stone"
    table.cell(0, 2).text = "The report must be retained for seven years."
    payload = io.BytesIO()
    document.save(payload)

    utterance = import_transcript(payload.getvalue(), "teams.docx").utterances[0]
    assert utterance.t0 == 135.0
    assert utterance.text == "Avery Stone: The report must be retained for seven years."


@pytest.mark.parametrize(
    ("filename", "status"),
    [("meeting.pdf", 415), ("../meeting.txt", 400), ("folder\\meeting.vtt", 400)],
)
def test_type_and_path_safety(filename: str, status: int) -> None:
    with pytest.raises(TranscriptImportError) as error:
        import_transcript("hello", filename)
    assert error.value.status_code == status


def test_size_limit_is_enforced_for_text_and_binary() -> None:
    with pytest.raises(TranscriptImportError) as error:
        import_transcript(b"x" * (MAX_IMPORT_BYTES + 1), "large.txt")
    assert error.value.status_code == 413


def test_invalid_and_expanding_docx_are_rejected() -> None:
    with pytest.raises(TranscriptImportError, match="valid DOCX"):
        import_transcript(b"not a zip", "fake.docx")

    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as output:
        output.writestr("word/document.xml", b"x" * (33 * 1024 * 1024))
    with pytest.raises(TranscriptImportError) as error:
        import_transcript(archive.getvalue(), "expanding.docx")
    assert error.value.status_code == 413


def test_empty_or_metadata_only_transcript_is_rejected() -> None:
    with pytest.raises(TranscriptImportError, match="empty"):
        import_transcript("", "empty.txt")
    with pytest.raises(TranscriptImportError, match="no readable utterances"):
        import_transcript("WEBVTT\n\nNOTE nothing to import", "empty.vtt")
